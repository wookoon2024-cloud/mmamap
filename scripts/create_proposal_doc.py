import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from pathlib import Path

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_borders(table, color="A0AEC0", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def add_caption(p, text):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = '맑은 고딕'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(90, 100, 120)
    r.italic = True

def create_document():
    dest_zip = Path(__file__).resolve().parent / "default_template.docx"
    doc = docx.Document(str(dest_zip))
    
    # Page Setup (A4, 20mm margins)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    
    # Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = '맑은 고딕'
    normal_style.font.size = Pt(9.5)
    normal_style.font.color.rgb = RGBColor(30, 30, 30)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(2)
    normal_style.paragraph_format.space_before = Pt(0)
    
    # Header tag [붙임 1]
    p_tag = doc.add_paragraph()
    r_tag = p_tag.add_run("【붙임 1】")
    r_tag.bold = True
    r_tag.font.size = Pt(11)
    r_tag.font.color.rgb = RGBColor(30, 60, 120)
    p_tag.paragraph_format.space_after = Pt(2)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("AI 활용 업무혁신 아이디어(혁신과제) 제안서")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(10, 30, 80)
    p_title.paragraph_format.space_after = Pt(10)
    
    # Create Table
    table = doc.add_table(rows=13, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="A0AEC0", sz="4")
    
    col_widths = [Inches(1.3), Inches(2.2), Inches(1.2), Inches(2.3)]
    
    HEADER_BG = "EEF2F7"
    SUBHEADER_BG = "E2E8F0"
    
    def format_cell(cell, text_runs, bg_color=None, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size=9.0):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        if bg_color:
            set_cell_shading(cell, bg_color)
        
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        
        if isinstance(text_runs, str):
            text_runs = [text_runs]
            
        for i, item in enumerate(text_runs):
            if i > 0:
                p = cell.add_paragraph()
                p.alignment = align
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                
            if isinstance(item, tuple):
                txt, is_b, color_rgb = item
                r = p.add_run(txt)
                r.bold = is_b
                r.font.name = '맑은 고딕'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                r.font.size = Pt(font_size)
                if color_rgb:
                    r.font.color.rgb = color_rgb
            else:
                r = p.add_run(item)
                r.bold = bold
                r.font.name = '맑은 고딕'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                r.font.size = Pt(font_size)
        return cell

    # Row 0: 과제명
    table.cell(0, 0).merge(table.cell(0, 0))
    table.cell(0, 1).merge(table.cell(0, 3))
    format_cell(table.cell(0, 0), "과 제 명", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(0, 1), [
        ("병역이행자와 소상공인이 함께하는 『군필지도』 상생 플랫폼 구축", True, RGBColor(10, 40, 110)),
        (" - 맞춤형 홍보물 자동 출력 지원을 통한 나라사랑가게 모집·이용 활성화 -", False, RGBColor(80, 80, 80))
    ], font_size=10.0)

    # Row 1: 제안 배경 및 필요성
    table.cell(1, 0).merge(table.cell(1, 0))
    table.cell(1, 1).merge(table.cell(1, 3))
    format_cell(table.cell(1, 0), "제안 배경 및\n필요성", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    cell_bg = table.cell(1, 1)
    format_cell(cell_bg, [
        ("□ 전국 단위 온라인 검색 중심 안내의 한계 및 현장 안내 수단 부족", True, None),
        ("  ○ 병무청 누리집 등 온라인 텍스트 검색 중심으로 운영되어, 병역이행자가 사전에 직접 검색하지 않는 한 혜택 인지 곤란", False, None),
        ("  ○ 실제 오프라인 매장 방문 시 해당 업소의 구체적 할인 혜택, 대상자(병역명문가, 복무자 등), 이용 방법 및 인근 연계 가맹점 정보를 확인할 수 있는 현장 안내 수단 부족", False, None),
        ("□ 가맹점(소상공인) 마케팅 지원 부재 및 개별 홍보물 제작 한계", True, None),
        ("  ○ 소상공인 사업자는 나라사랑가게에 참여해도 실질적 매장 홍보 도구가 없어 자발적 참여 유인 부족", False, None),
        ("  ○ 전국 가맹점마다 개별 맞춤형 홍보물을 수작업으로 디자인·인쇄 배부하기에는 행정 여건상 곤란", False, None),
        ("□ 온·오프라인 연계 및 가맹점 자율 활용 플랫폼 구축 필요", True, None),
        ("  ○ 해당 매장을 중심으로 주변 가맹점을 함께 소개하는 맞춤형 팜플렛을 자동 생성하고, 가맹점 사업자가 사이트에서 직접 확인·출력하여 매장에 비치·활용할 수 있는 지속 가능한 상생 체계 마련 시급", False, None),
    ])
    
    # 1. 이미지: AS-IS vs TO-BE 좌우 2열 배치
    img_asis = Path("img/screenshot_asis_homepage.png")
    img_tobe = Path("img/screenshot_tobe_map.png")
    if img_asis.exists() and img_tobe.exists():
        p_img_tbl = cell_bg.add_paragraph()
        p_img_tbl.paragraph_format.space_before = Pt(4)
        p_img_tbl.paragraph_format.space_after = Pt(2)
        
        # Add 2-column inner table for side-by-side comparison
        sub_tbl = cell_bg.add_table(rows=2, cols=2)
        sub_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(sub_tbl, color="CBD5E0", sz="2")
        
        # Images
        c1 = sub_tbl.cell(0, 0)
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c1.paragraphs[0].add_run().add_picture(str(img_asis), width=Inches(2.5))
        
        c2 = sub_tbl.cell(0, 1)
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c2.paragraphs[0].add_run().add_picture(str(img_tobe), width=Inches(2.5))
        
        # Captions
        c1_cap = sub_tbl.cell(1, 0)
        add_caption(c1_cap.paragraphs[0], "[현행(AS-IS)] 누리집 텍스트/표 중심 안내\n(위치 파악 곤란 및 현장 체감도 저하)")
        
        c2_cap = sub_tbl.cell(1, 1)
        add_caption(c2_cap.paragraphs[0], "[개선(TO-BE)] 『군필지도』 전국 GIS 시각화\n(한눈에 보는 혜택 탐색 및 상생 연계)")

    # Row 2: 연구분야 & 소속기관
    format_cell(table.cell(2, 0), "연구분야", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(2, 1), "동원·소집 / 고객지원 / 정보기획", align=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(table.cell(2, 2), "소속기관", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(2, 3), "병무청 본청 (또는 OO지방병무청)", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 3: 대표직원(1인)
    format_cell(table.cell(3, 0), "대표직원(1인)", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(3, 1), "소속: 동원관리과", align=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(table.cell(3, 2), "직급/성명", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(3, 3), "행정주사 O O O", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 4: 참여인원
    format_cell(table.cell(4, 0), "참여인원", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(4, 1), "소속: 고객지원과, 정보기획과", align=WD_ALIGN_PARAGRAPH.CENTER)
    format_cell(table.cell(4, 2), "직급/성명", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(4, 3), "행정주사 O O O, 행정서기 O O O", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 5: 구분 헤더
    table.cell(5, 0).merge(table.cell(5, 3))
    format_cell(table.cell(5, 0), "과 제 개 요  및  추 진 내 용", bg_color=SUBHEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=10.0)

    # Row 6: 워크메이트 활용 (가점-5점)
    table.cell(6, 0).merge(table.cell(6, 0))
    table.cell(6, 1).merge(table.cell(6, 3))
    format_cell(table.cell(6, 0), "워크메이트 활용\n(가점-5점)", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(6, 1), [
        ("□ [분석·통계 에이전트]를 활용한 지역별 우대 인프라 및 필요 업종 분석", True, None),
        ("  ○ 전국 14개 지방청 우대 DB 통계 분석: 전국 우대기관 데이터와 지역별 청년 인구 분포를 분석하여 지자체별 혜택 인프라 현황 도출", False, None),
        ("  ○ 청년 선호 필요·부족 업종 분석: 요식업 외에 병역이행자 선호 업종(스터디카페, 헬스장, 숙박, 의료 등)의 지역별 부족 현황 분석", False, None),
        ("  ○ 제도 운영 기초자료 활용: 데이터 분석 결과를 바탕으로 향후 나라사랑가게 제도 운영 및 정책 수립 기초자료로 활용", False, None),
    ])

    # Row 7: 활용 모델
    table.cell(7, 0).merge(table.cell(7, 0))
    table.cell(7, 1).merge(table.cell(7, 3))
    format_cell(table.cell(7, 0), "활용 모델", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(7, 1), [
        ("□ 제미나이(Gemini), 구글 플로우(Google Flow), AI 워크메이트 복합 활용", True, None),
        ("  ○ 제미나이 (Gemini): 가맹점별 맞춤형 상생 홍보 문구 생성 및 대상별 혜택 요약 텍스트 정제", False, None),
        ("  ○ 구글 플로우 (Google Flow): 팜플렛·홍보물에 적용될 비주얼 그래픽 및 템플릿 디자인 에셋 생성", False, None),
        ("  ○ AI 워크메이트 (분석·통계 에이전트): 전국 우대기관 데이터 통계 분석 및 필요 업종 현황 분석", False, None),
    ])

    # Row 8: 구현 방법
    table.cell(8, 0).merge(table.cell(8, 0))
    table.cell(8, 1).merge(table.cell(8, 3))
    format_cell(table.cell(8, 0), "구현 방법", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    cell_impl = table.cell(8, 1)
    format_cell(cell_impl, [
        ("□ 『군필지도』 웹 GIS 플랫폼 및 다형태 원클릭 인쇄 엔진 구축", True, None),
        ("  ○ 매장 중심 맞춤형 상생지도 자동 제작: 특정 가맹점 선택 시, 해당 매장을 중심점으로 인근 가맹점 위치 및 혜택 정보를 결합한 ‘나라사랑 상생지도’ 자동 렌더링", False, None),
        ("  ○ 업종·매장 환경별 다양한 홍보물 규격 지원: A4 포스터/리플렛형, 미니 테이블 스탠드형, 도어행거(문고리)형 등 맞춤 규격 지원", False, None),
        ("  ○ 매장별 고유 QR코드 기반 활성화 분석 (핵심 기능): 각 팜플렛마다 가맹점별 고유 QR코드를 자동 삽입하여 스캔 로그를 집계·분석하고 제도 개선 피드백 데이터로 활용", False, None),
        ("  ○ 가맹점 사업자 자율 활용 체계: 사업자가 사이트에 접속하여 언제든지 매장 맞춤형 팜플렛을 확인하고 직접 자율 출력하여 매장에 비치·홍보", False, None),
    ])
    
    # 2. 이미지: 구현 화면
    if img_tobe.exists():
        p_img_impl = cell_impl.add_paragraph()
        p_img_impl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img_impl.paragraph_format.space_before = Pt(4)
        p_img_impl.paragraph_format.space_after = Pt(2)
        p_img_impl.add_run().add_picture(str(img_tobe), width=Inches(4.8))
        add_caption(cell_impl.add_paragraph(), "[그림 1] 실제 구현 완료된 『군필지도』 웹 GIS 플랫폼 메인 구동 화면 (전국 가맹점 연동)")

    # Row 9: 데이터활용
    table.cell(9, 0).merge(table.cell(9, 0))
    table.cell(9, 1).merge(table.cell(9, 3))
    format_cell(table.cell(9, 0), "데이터활용", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(9, 1), [
        ("□ 공공데이터, 네이버 지도 API 및 자체 공간 DB 활용", True, None),
        ("  ○ 병무청 공공데이터: 전국 14개 지방병무청 우대기관 Open API 및 JSON 데이터", False, None),
        ("  ○ 네이버 지도 API & 자체 GIS DB: Naver Maps Geocoding을 통한 좌표 변환 및 반경 검색용 자체 공간 데이터베이스 구축", False, None),
        ("  ○ 비식별 QR 스캔 통계: 매장별 스캔 횟수 및 시간대 로그 분석 (개인정보 배제 및 보안 지침 준수)", False, None),
    ])

    # Row 10: 세부 개선 내용 (AS-IS)
    table.cell(10, 0).merge(table.cell(11, 0))
    format_cell(table.cell(10, 0), "세부 개선 내용", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    format_cell(table.cell(10, 1), "현 행\n(AS-IS)", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    table.cell(10, 2).merge(table.cell(10, 3))
    format_cell(table.cell(10, 2), [
        ("□ 한눈에 보는 지도 부재 및 현장 안내 한계: 텍스트 나열식으로 위치 파악 곤란, 매장 방문 시 혜택 확인 불가", False, None),
        ("□ 경험 의존적 관리 및 이용 현황 파악 불가: 통계 도구 부재로 부족 업종 파악 미흡, 협약 후 실제 이용 여부 모니터링 불가", False, None),
    ])

    # Row 11: 세부 개선 내용 (TO-BE)
    format_cell(table.cell(11, 1), "개 선\n(TO-BE)", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    table.cell(11, 2).merge(table.cell(11, 3))
    cell_tobe = table.cell(11, 2)
    format_cell(cell_tobe, [
        ("□ 전국 『군필지도』 및 맞춤형 상생 팜플렛 제공: 전국 나라사랑가게 지도 검색 및 매장 중심 상생 팜플렛 즉시 생성", False, None),
        ("□ 다양한 템플릿 자율 출력 & QR 스캔 분석: A4/스탠드형 등 사업자 직접 출력, QR 스캔 통계 기반 활성화 분석 및 정책 환류", False, None),
    ])
    
    # 3. 이미지: 프로세스 비교 인포그래픽
    img_proc = Path("img/infographic_process.jpg")
    if img_proc.exists():
        p_img_proc = cell_tobe.add_paragraph()
        p_img_proc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img_proc.paragraph_format.space_before = Pt(4)
        p_img_proc.paragraph_format.space_after = Pt(2)
        p_img_proc.add_run().add_picture(str(img_proc), width=Inches(4.5))
        add_caption(cell_tobe.add_paragraph(), "[그림 2] 나라사랑가게 운영 프로세스 혁신 (AS-IS vs TO-BE 구조 비교 인포그래픽)")

    # Row 12: 기대 효과
    table.cell(12, 0).merge(table.cell(12, 0))
    table.cell(12, 1).merge(table.cell(12, 3))
    format_cell(table.cell(12, 0), "기대 효과", bg_color=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    cell_eff = table.cell(12, 1)
    format_cell(cell_eff, [
        ("□ 정량적 기대효과 (데이터 및 관리 효율)", True, None),
        ("  ○ 가맹점별 방문자 수 및 활성화 통계 데이터 확보: 매장별 QR 스캔 데이터를 통해 실제 이용 실적(스캔 건수, 시간대, 지역별 접속량)을 실시간 수치로 측정 및 모니터링", False, None),
        ("  ○ 데이터 기반 [업종별 맞춤 활성화 방안] 수립: 업종별 QR 이용 통계 분석을 통해 이용률 저조 업종 보완 및 인기 업종 확충 근거 마련", False, None),
        ("  ○ 전국 14개 지방청 우대 데이터 일원화: 14개 청별로 분산 관리되던 가맹점 DB를 단일 플랫폼으로 통합 관리 (14개 분산 ➔ 1개 통합 DB)", False, None),
        ("  ○ 정보 탐색 및 팜플렛 생성 소요시간 단축: 텍스트 목록 검색 ➔ 지도 기반 원클릭 즉시 탐색 및 맞춤 팜플렛 즉시 생성", False, None),
        ("□ 정성적 기대효과 (대국민 체감 및 정책 가치)", True, None),
        ("  ○ 전국 단위 우대 인프라 현황 총괄 조망 및 지역 간 단절 해소: 지방청별 칸막이를 없애고 전국 나라사랑가게 현황을 한눈에 직관적으로 파악", False, None),
        ("  ○ 병역의무자의 혜택 탐색 및 방문 편의 극대화: 전국 나라사랑가게를 지도로 시각화하여 내 주변 혜택 매장을 쉽게 찾고 직접 방문 가능", False, None),
        ("  ○ 소상공인의 주변 제휴업체 정보 파악 및 상생 연계: 매장 주변의 다른 나라사랑가게를 지도에서 확인하고 상생 팜플렛을 통해 상호 홍보", False, None),
        ("  ○ 신규 소상공인 참여 유입 활성화: 지도 노출 및 홍보 도구 지원으로 인근 소상공인의 자발적 가맹 신청 유도", False, None),
    ])
    
    # 4. 이미지: 매장 비치 목업 & QR 통계 화면
    img_mockup = Path("img/mockup_table_qr.jpg")
    if img_mockup.exists():
        p_img_eff = cell_eff.add_paragraph()
        p_img_eff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img_eff.paragraph_format.space_before = Pt(4)
        p_img_eff.paragraph_format.space_after = Pt(2)
        p_img_eff.add_run().add_picture(str(img_mockup), width=Inches(4.8))
        add_caption(cell_eff.add_paragraph(), "[그림 3] 가맹점 매장 내 맞춤형 팜플렛 비치 예시 및 실시간 QR코드 스캔 통계 대시보드")

    # Adjust widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            if idx < len(row.cells):
                row.cells[idx].width = width

    out_path = Path("AI_활용_업무혁신_아이디어_제안서_군필지도.docx").resolve()
    doc.save(str(out_path))
    print(f"File updated with embedded images: {out_path}")
    return str(out_path)

if __name__ == "__main__":
    create_document()
